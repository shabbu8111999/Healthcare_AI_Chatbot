"""
Builds a downloadable PDF or Word file from a chat history.

Kept separate from app.py so the export formatting logic has its own
home, same reasoning as chat_store.py being its own file.
"""

import io

from docx import Document as DocxDocument
from fpdf import FPDF


def build_pdf_bytes(history: list) -> bytes:
    pdf = FPDF()
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Healthcare Assistant Chat", ln=True)
    # multi_cell and cell with width 0 leave the x position near the
    # right edge of the page instead of resetting it, without this the
    # very next multi_cell call has almost no width left to work with
    # and fpdf raises "not enough horizontal space"
    pdf.set_x(pdf.l_margin)
    pdf.ln(2)

    for message in history:
        speaker = "You" if message["role"] == "user" else "Assistant"

        pdf.set_font("Helvetica", "B", 11)
        pdf.multi_cell(0, 7, speaker + ":")
        pdf.set_x(pdf.l_margin)

        pdf.set_font("Helvetica", size=11)
        # the base PDF fonts only support latin-1, so anything outside
        # that gets swapped for a safe substitute instead of crashing
        safe_text = message["content"].encode("latin-1", errors="replace").decode("latin-1")
        pdf.multi_cell(0, 7, safe_text)
        pdf.set_x(pdf.l_margin)
        pdf.ln(3)

    return bytes(pdf.output())


def build_docx_bytes(history: list) -> bytes:
    doc = DocxDocument()
    doc.add_heading("Healthcare Assistant Chat", level=1)

    for message in history:
        speaker = "You" if message["role"] == "user" else "Assistant"
        speaker_paragraph = doc.add_paragraph()
        speaker_paragraph.add_run(speaker + ":").bold = True
        doc.add_paragraph(message["content"])

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()